"""Генерация .docx документа стратегии личного бренда из данных опроса."""

import io
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _split_lines(text):
    """Разбивает текст на строки, удаляя пустые."""
    if not text:
        return []
    return [line.strip() for line in text.strip().split("\n") if line.strip()]


def _split_pipe(text):
    """Разбивает строки по символу '|' на колонки."""
    rows = []
    for line in _split_lines(text):
        cols = [c.strip() for c in line.split("|")]
        rows.append(cols)
    return rows


def _add_table(doc, headers, rows):
    """Добавляет таблицу с заголовками и строками."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row_data in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_data):
            cells[i].text = val if i < len(row_data) else ""
    doc.add_paragraph()


def generate_brand_docx(data: dict) -> bytes:
    """Генерирует .docx документ стратегии личного бренда или организации.
    data — dict с полями опроса.
    Возвращает байты документа."""
    doc = Document()

    brand_type = data.get("brand_type", "personal")
    is_org = (brand_type == "organization")
    is_community = (brand_type == "community")

    # === ТИТУЛ ===
    if is_community:
        entity_name = data.get("org_name", "")
        title = doc.add_heading(f"Стратегия развития бизнес-сообщества\n{entity_name}", 0)
    elif is_org:
        entity_name = data.get("org_name", "")
        title = doc.add_heading(f"Стратегия бренд-коммуникаций\n{entity_name}", 0)
    else:
        entity_name = data.get("client_name", "")
        title = doc.add_heading(f"Стратегия развития личного бренда\n{entity_name}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    version = data.get("version", "1.0")
    date_str = data.get("date_str") or datetime.now().strftime("%B %Y")
    doc.add_paragraph(f"Версия {version}  {date_str}")
    doc.add_paragraph()

    # === 1. РЕЗЮМЕ И ЦЕЛИ ===
    doc.add_heading("1. Резюме и цели", level=1)

    if is_community:
        doc.add_heading("Общественная организация:", level=3)
        doc.add_paragraph(f"Название: {data.get('org_name', '—')}")
        doc.add_paragraph(f"Сфера деятельности: {data.get('org_industry', '—')}")
        doc.add_paragraph(f"Миссия: {data.get('org_mission', '—')}")
        doc.add_paragraph(f"Направления деятельности: {data.get('org_products', '—')}")
        doc.add_paragraph(f"Участников: {data.get('org_team_size', '—')}")
        doc.add_paragraph(f"Год основания: {data.get('org_founded', '—')}")
        doc.add_paragraph(f"Формат участия: {data.get('org_format', '—')}")
    elif is_org:
        doc.add_heading("Организация:", level=3)
        doc.add_paragraph(f"Название: {data.get('org_name', '—')}")
        doc.add_paragraph(f"Сфера деятельности: {data.get('org_industry', '—')}")
        doc.add_paragraph(f"Миссия: {data.get('org_mission', '—')}")
        doc.add_paragraph(f"Продукты/услуги: {data.get('org_products', '—')}")
        doc.add_paragraph(f"Сотрудников: {data.get('org_team_size', '—')}")
        doc.add_paragraph(f"Год основания: {data.get('org_founded', '—')}")
    else:
        doc.add_heading("Текущая ситуация:", level=3)
        doc.add_paragraph(data.get("current_situation", "—"))

    doc.add_heading("Ключевые цели на 12 месяцев:", level=3)

    doc.add_paragraph(
        f"Качественный рост аудитории: Набрать {data.get('goal_audience_number', '—')} "
        f"подписчиков в основном {data.get('goal_audience_platform', '—')} "
        f"(активная аудитория: {data.get('goal_audience_description', '—')})."
    )
    doc.add_paragraph(
        f"Укрепление экспертного статуса: Стать признанным федеральным экспертом "
        f"в сфере {data.get('goal_expert_field', '—')} "
        f"(не менее {data.get('goal_expert_speaking', '—')} приглашений выступить "
        f"на федеральных площадках в год)."
    )
    doc.add_paragraph(
        f"Коммерциализация личного бренда: Обеспечить поток из не менее "
        f"{data.get('goal_commercial_leads', '—')} заявок в месяц на платные услуги "
        f"({data.get('goal_commercial_services', '—')}) через личный бренд."
    )
    doc.add_paragraph(
        f"Системный контент: {data.get('goal_content_system', '—')}"
    )

    # === 2. ПОЗИЦИОНИРОВАНИЕ И УТП ===
    doc.add_heading("2. Позиционирование и УТП", level=1)

    doc.add_heading("Сущность бренда:", level=3)
    doc.add_paragraph(data.get("brand_essence", "—"))

    doc.add_heading("Позиционирование (Ответы на ключевые вопросы):", level=3)
    if is_community:
        doc.add_paragraph(f"Для кого создано: {data.get('positioning_who', '—')}")
        doc.add_paragraph(f"Какую проблему решает: {data.get('positioning_problem', '—')}")
        doc.add_paragraph(f"Ценность для участника: {data.get('positioning_result', '—')}")
        doc.add_paragraph(f"Чем отличаемся от других: {data.get('positioning_difference', '—')}")
        doc.add_paragraph(f"Форматы взаимодействия: {data.get('community_formats', '—')}")
    elif is_org:
        doc.add_paragraph(f"Кому помогаем: {data.get('positioning_who', '—')}")
        doc.add_paragraph(f"Какую проблему решаем: {data.get('positioning_problem', '—')}")
        doc.add_paragraph(f"Какой результат получает клиент: {data.get('positioning_result', '—')}")
        doc.add_paragraph(f"Почему выбирают нас: {data.get('positioning_difference', '—')}")
    else:
        doc.add_paragraph(f"Кому я помогаю? {data.get('positioning_who', '—')}")
        doc.add_paragraph(f"Какую проблему решаю? {data.get('positioning_problem', '—')}")
        doc.add_paragraph(f"Какой результат получает клиент? {data.get('positioning_result', '—')}")

    doc.add_heading("Уникальное торговое предложение (УТП):", level=3)
    doc.add_paragraph(data.get("usp", "—"))

    # === 3. АУДИТОРИИ И КАНАЛЫ ===
    doc.add_heading("3. Аудитории и каналы коммуникации", level=1)

    doc.add_heading("Идеальные подписчики/клиенты:", level=3)
    for line in _split_lines(data.get("ideal_audiences", "")):
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("Карта каналов (по приоритету):", level=3)
    channels = data.get("channels", [])
    if channels:
        _add_table(doc, ["Канал", "Тип", "Частота", "Описание"], [
            [ch.get("platform", ""), ch.get("type", ""), ch.get("frequency", ""), ch.get("description", "")]
            for ch in channels
        ])
    else:
        doc.add_paragraph("—")

    # === 4. ВИЗУАЛЬНАЯ И ВЕРБАЛЬНАЯ ИДЕНТИЧНОСТЬ ===
    doc.add_heading("4. Визуальная и вербальная идентичность", level=1)

    doc.add_heading("А. Визуальная идентичность", level=3)
    doc.add_paragraph(f"Концепция: {data.get('visual_concept', '—')}")

    doc.add_paragraph("Цветовая палитра:")
    _add_table(doc, ["Роль", "Название", "HEX"], [
        ["Основной", data.get("color_main_name", ""), data.get("color_main_hex", "")],
        ["Акцент 1", data.get("color_accent1_name", ""), data.get("color_accent1_hex", "")],
        ["Акцент 2", data.get("color_accent2_name", ""), data.get("color_accent2_hex", "")],
        ["Нейтральная база", data.get("color_neutral", ""), ""],
    ])

    doc.add_paragraph("Типографика:")
    doc.add_paragraph(f"Заголовки: {data.get('typography_headings', '—')}", style="List Bullet")
    doc.add_paragraph(f"Текст: {data.get('typography_body', '—')}", style="List Bullet")
    doc.add_paragraph(f"Акценты: {data.get('typography_accents', '—')}", style="List Bullet")

    doc.add_heading("Б. Tone of Voice (Тон коммуникации)", level=3)
    doc.add_paragraph(data.get("tov_description", "—"))
    doc.add_paragraph(f"Референсы: {data.get('tov_references', '—')}")

    tov_axes = _split_lines(data.get("tov_axes", ""))
    if tov_axes:
        doc.add_paragraph("Оси ToV:")
        for axis in tov_axes:
            doc.add_paragraph(axis, style="List Number")

    # === 5. КОНТЕНТ-СТРАТЕГИЯ ===
    doc.add_heading("5. Контент-стратегия: Три столпа", level=1)
    doc.add_paragraph(
        "Контент строится на балансе трех направлений, работающих на разные этапы воронки доверия."
    )

    pillars = [
        ("1", data.get("pillar1_name", "Кто я"), data.get("pillar1_goal", ""),
         data.get("pillar1_formats", ""), data.get("pillar1_frequency", "30%")),
        ("2", data.get("pillar2_name", "Решаю боль"), data.get("pillar2_goal", ""),
         data.get("pillar2_formats", ""), data.get("pillar2_frequency", "50%")),
        ("3", data.get("pillar3_name", "Втягиваю"), data.get("pillar3_goal", ""),
         data.get("pillar3_formats", ""), data.get("pillar3_frequency", "20%")),
    ]

    _add_table(doc, ["Столп", "Цель", "Ключевые форматы", "Частота"], [
        [f"{p[1]}", p[2], p[3], p[4]]
        for p in pillars
    ])

    special = data.get("special_formats", "")
    if special:
        doc.add_paragraph(f"Особые форматы: {special}")

    # === 6. ВОРОНКА И ПРОДУКТЫ ===
    doc.add_heading("6. Воронка вовлечения и продуктовый ряд", level=1)

    doc.add_heading("Воронка (путь от незнакомца к клиенту):", level=3)
    funnel_stages = data.get("funnel", [])
    if funnel_stages:
        _add_table(doc, ["Этап", "Контент", "CTA"], [
            [s.get("name", ""), s.get("content", ""), s.get("cta", "")]
            for s in funnel_stages
        ])
    else:
        doc.add_paragraph("—")

    doc.add_heading("Продуктовый ряд (от простого к сложному):", level=3)
    doc.add_paragraph(f"Входная точка (Lead Magnet): {data.get('product_lead_magnet', '—')}")
    doc.add_paragraph(f"Самая простая услуга: {data.get('product_simple', '—')}")
    doc.add_paragraph(f"Услуга-конвертер: {data.get('product_converter', '—')}")
    doc.add_paragraph(f"Комплексные решения: {data.get('product_complex', '—')}")
    doc.add_paragraph(f"Экспертные продукты: {data.get('product_expert', '—')}")

    # === 7. МЕТРИКИ ===
    doc.add_heading("7. Метрики", level=1)

    quant_rows = _split_pipe(data.get("metrics_quantitative", ""))
    if quant_rows:
        doc.add_heading("Количественные:", level=3)
        _add_table(doc, ["Метрика", "Целевое значение", "Измерение"], quant_rows)

    qual_rows = _split_pipe(data.get("metrics_qualitative", ""))
    if qual_rows:
        doc.add_heading("Качественные:", level=3)
        _add_table(doc, ["Метрика", "Целевое значение", "Измерение"], qual_rows)

    comm_rows = _split_pipe(data.get("metrics_commercial", ""))
    if comm_rows:
        doc.add_heading("Коммерческие:", level=3)
        _add_table(doc, ["Метрика", "Целевое значение", "Измерение"], comm_rows)

    # === 8. РИСКИ И РЕКОМЕНДАЦИИ ===
    doc.add_heading("8. Потенциальные риски и рекомендации", level=1)

    risk_rows = _split_pipe(data.get("risks", ""))
    if risk_rows:
        doc.add_heading("Потенциальные риски:", level=3)
        _add_table(doc, ["Риск", "Решение"], risk_rows)

    recs = _split_lines(data.get("recommendations", ""))
    if recs:
        doc.add_heading("Рекомендации:", level=3)
        for rec in recs:
            doc.add_paragraph(rec, style="List Bullet")

    # === 9. АРХЕТИПИЧЕСКИЙ ОБРАЗ ===
    doc.add_heading("9. Архетипический образ", level=1)

    doc.add_heading(f"Ключевой архетип: {data.get('archetype_key_name', '—')}", level=3)
    doc.add_paragraph(f"Суть: {data.get('archetype_key_essence', '—')}")
    manifestations = _split_lines(data.get("archetype_key_manifestations", ""))
    if manifestations:
        doc.add_paragraph("Проявления:")
        for m in manifestations:
            doc.add_paragraph(m, style="List Bullet")

    supporting = _split_pipe(data.get("archetype_supporting", ""))
    if supporting:
        doc.add_heading("Вспомогательные архетипы:", level=3)
        _add_table(doc, ["Архетип", "Суть", "Проявления"], supporting)

    doc.add_heading("Синтез архетипов:", level=3)
    doc.add_paragraph(data.get("archetype_synthesis", "—"))

    table_rows = _split_pipe(data.get("archetype_table", ""))
    if table_rows:
        doc.add_heading("Таблица проявлений:", level=3)
        _add_table(doc, ["Ситуация / Контент-формат", "Доминирующий архетип", "Как это выглядит и звучит"], table_rows)

    # === 10. ПРОМПТ ДЛЯ ГЕНЕРАЦИИ КОНТЕНТА ===
    doc.add_heading("10. Промпт для генерации контента", level=1)

    doc.add_heading("Контекст и личность:", level=3)
    doc.add_paragraph(data.get("prompt_context", "—"))

    doc.add_heading("Tone of Voice:", level=3)
    doc.add_paragraph(data.get("prompt_tov", "—"))

    doc.add_heading("Структура мысли:", level=3)
    doc.add_paragraph(data.get("prompt_structure", "—"))

    doc.add_heading("Контент-направления:", level=3)
    for line in _split_lines(data.get("prompt_directions", "")):
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("Форматы и шаблоны:", level=3)
    doc.add_paragraph(data.get("prompt_formats", "—"))

    doc.add_heading("Технические требования к тексту:", level=3)
    doc.add_paragraph(data.get("prompt_technical", "—"))

    # === Сериализация в байты ===
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
