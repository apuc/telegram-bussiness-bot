from pathlib import Path
import json
import os
import re
import threading
import uuid

import bcrypt
from fastapi import Depends, FastAPI, Form, HTTPException, Query, Request
from fastapi.responses import JSONResponse, RedirectResponse, Response
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from starlette.middleware.sessions import SessionMiddleware

from ai_core.image_generator import IMAGE_SIZES, IMAGE_TEMPLATES, STYLE_PRESETS, ImageGenerator, detect_ext, get_sizes_for_model
from ai_core.neural_network import ContentGenerator
from ai_core.social_analyzer import analyze_posts_with_ai, detect_platform, fetch_telegram_posts
from ai_core.telegram_publisher import (
    check_bot_is_admin,
    publish_post,
    resolve_chat,
)
from cabinet_service.brand_doc_generator import generate_brand_docx
from config import SECRET_KEY
from database import (
    add_post_revision,
    create_task,
    create_web_user,
    delete_brand_survey,
    delete_channel,
    delete_media_by_id,
    delete_plan_item,
    delete_post_by_id,
    delete_post_image,
    delete_user_by_id,
    get_brand_survey_by_id,
    get_brand_surveys_for_user,
    get_channel_by_id,
    get_channels_for_user,
    get_image_provider_settings,
    get_latest_brand_survey_for_user,
    get_media_by_id_for_user,
    get_media_for_user,
    get_plan_for_user,
    get_plan_stats_for_user,
    get_post_by_id_for_user,
    get_post_images,
    get_post_revisions,
    get_posts_for_user,
    get_provider_settings,
    get_stats_for_user,
    get_task,
    get_text_provider_settings,
    get_user_by_email,
    get_user_by_id,
    is_channel_already_added,
    mark_post_published,
    PROVIDER_PRESETS,
    replace_plan_for_user,
    save_brand_survey,
    save_channel,
    save_media,
    save_post_for_user,
    save_post_image,
    save_provider_settings,
    update_plan_status,
    update_post_content,
    update_task_status,
    update_user_field_by_id,
)

BASE_DIR = Path(__file__).resolve().parent
MEDIA_DIR = BASE_DIR / "media"
MEDIA_DIR.mkdir(exist_ok=True)

app = FastAPI(title="Личный кабинет")
app.add_middleware(SessionMiddleware, secret_key=SECRET_KEY)
app.mount("/static", StaticFiles(directory=str(BASE_DIR / "static")), name="static")
app.mount("/media", StaticFiles(directory=str(MEDIA_DIR)), name="media")
templates = Jinja2Templates(directory=str(BASE_DIR / "templates"))

content_gen = ContentGenerator()
image_gen = ImageGenerator()


def get_user_content_gen(user: dict) -> ContentGenerator:
    """Возвращает ContentGenerator с настройками text-провайдера пользователя."""
    settings = get_text_provider_settings(user["id"])
    if settings and settings.get("api_key"):
        return ContentGenerator(provider_settings=settings)
    return content_gen


def get_user_image_gen(user: dict) -> ImageGenerator:
    """Возвращает ImageGenerator с настройками image-провайдера пользователя."""
    settings = get_image_provider_settings(user["id"])
    if settings and settings.get("api_key"):
        return ImageGenerator(provider_settings=settings)
    return image_gen

PLATFORMS = [
    ("instagram", "Instagram"),
    ("linkedin", "LinkedIn"),
    ("telegram", "Telegram"),
    ("twitter", "Twitter/X"),
    ("ad", "Рекламный пост"),
]

PLATFORM_LABELS = dict(PLATFORMS)

def platform_label(platform: str) -> str:
    return PLATFORM_LABELS.get(platform, platform)


class NotAuthenticated(Exception):
    pass


@app.exception_handler(NotAuthenticated)
async def not_authenticated_handler(request: Request, exc: NotAuthenticated):
    return RedirectResponse(url="/login", status_code=303)


def current_user(request: Request) -> dict:
    user_id = request.session.get("user_id")
    if not user_id:
        raise NotAuthenticated()
    user = get_user_by_id(user_id)
    if not user:
        request.session.clear()
        raise NotAuthenticated()
    return user


def flash(request: Request, message: str):
    request.session["flash"] = message


def render(request: Request, template: str, status_code: int = 200, **context):
    context.setdefault("user", None)
    context["flash_message"] = request.session.pop("flash", None)
    return templates.TemplateResponse(request, template, context, status_code=status_code)


def business_name_of(user: dict) -> str:
    """Возвращает название вида бизнеса. Оставлено для совместимости (профили без медиа-образа)."""
    from keyboards.reply_keyboards import BUSINESS_TYPES
    return BUSINESS_TYPES.get(str(user.get("business_type")), {}).get("name", "Не указан")


@app.get("/health")
def health():
    return {"status": "ok"}


@app.get("/")
def root(request: Request):
    # Посадочная страница с описанием сервиса
    return render(request, "landing.html")


# === РЕГИСТРАЦИЯ ===

@app.get("/register")
def register_form(request: Request):
    if request.session.get("user_id"):
        return RedirectResponse("/dashboard", status_code=303)
    return render(request, "register.html")


@app.post("/register")
def register_submit(
    request: Request,
    name: str = Form(...),
    phone: str = Form(...),
    email: str = Form(...),
    password: str = Form(...),
    password_confirm: str = Form(...),
):
    email = email.strip().lower()
    errors = []
    if len(password) < 8:
        errors.append("Пароль должен быть не короче 8 символов.")
    if password != password_confirm:
        errors.append("Пароли не совпадают.")
    if "@" not in email or "." not in email.split("@")[-1]:
        errors.append("Введи корректный email.")
    if not errors and get_user_by_email(email):
        errors.append("Пользователь с таким email уже зарегистрирован.")

    if errors:
        return render(
            request, "register.html", status_code=400,
            errors=errors,
            form={"name": name, "phone": phone, "email": email},
        )

    password_hash = bcrypt.hashpw(password.encode("utf-8"), bcrypt.gensalt()).decode("utf-8")
    user_id = create_web_user(email, password_hash, name, phone)
    request.session["user_id"] = user_id
    flash(request, "Регистрация завершена! Добро пожаловать.")
    return RedirectResponse("/dashboard", status_code=303)


# === ВХОД / ВЫХОД ===

@app.get("/login")
def login_form(request: Request):
    if request.session.get("user_id"):
        return RedirectResponse("/dashboard", status_code=303)
    return render(request, "login.html")


@app.post("/login")
def login_submit(request: Request, email: str = Form(...), password: str = Form(...)):
    email = email.strip().lower()
    user = get_user_by_email(email)
    valid = (
        user is not None
        and user["password_hash"]
        and bcrypt.checkpw(password.encode("utf-8"), user["password_hash"].encode("utf-8"))
    )
    if not valid:
        return render(
            request, "login.html", status_code=400,
            errors=["Неверный email или пароль."], form={"email": email},
        )

    request.session["user_id"] = user["id"]
    return RedirectResponse("/dashboard", status_code=303)


@app.post("/logout")
def logout(request: Request):
    request.session.clear()
    return RedirectResponse("/login", status_code=303)


# === КАБИНЕТ ===

@app.get("/dashboard")
def dashboard(request: Request, user: dict = Depends(current_user)):
    latest_survey = get_latest_brand_survey_for_user(user["id"])
    return render(request, "dashboard.html", user=user, latest_survey=latest_survey)


@app.get("/profile/edit")
def profile_edit_form(request: Request, user: dict = Depends(current_user)):
    provider = get_provider_settings(user["id"])
    return render(request, "profile_edit.html", user=user,
                  provider=provider, provider_presets=PROVIDER_PRESETS)


@app.post("/profile/edit")
def profile_edit_submit(
    request: Request,
    name: str = Form(...),
    phone: str = Form(...),
    user: dict = Depends(current_user),
):
    update_user_field_by_id(user["id"], "name", name)
    update_user_field_by_id(user["id"], "phone", phone)
    flash(request, "Профиль обновлён.")
    return RedirectResponse("/dashboard", status_code=303)


# === НАСТРОЙКИ ПРОВАЙДЕРА ИИ ===

@app.get("/provider")
def provider_form(request: Request, user: dict = Depends(current_user)):
    provider = get_provider_settings(user["id"])
    return render(request, "provider.html", user=user,
                  provider=provider, provider_presets=PROVIDER_PRESETS)


@app.post("/provider")
def provider_submit(
    request: Request,
    # Text provider
    text_provider_type: str = Form(...),
    text_api_key: str = Form(""),
    text_base_url: str = Form(""),
    text_model: str = Form(""),
    text_use_global: str = Form(""),
    # Image provider
    image_provider_type: str = Form(...),
    image_api_key: str = Form(""),
    image_base_url: str = Form(""),
    image_model: str = Form(""),
    image_use_global: str = Form(""),
    user: dict = Depends(current_user),
):
    errors = []

    # Валидация text-провайдера
    if text_provider_type not in PROVIDER_PRESETS:
        errors.append("Выберите провайдера для текста.")
    text_preset = PROVIDER_PRESETS.get(text_provider_type, {})
    if not text_use_global and text_preset.get("needs_key") and not text_api_key.strip():
        errors.append("Укажите API ключ для текстового провайдера или выберите «Использовать глобальные настройки».")
    if not text_use_global and text_provider_type == "custom" and not text_base_url.strip():
        errors.append("Для кастомного текстового провайдера укажите Base URL.")

    # Валидация image-провайдера
    if image_provider_type not in PROVIDER_PRESETS:
        errors.append("Выберите провайдера для изображений.")
    image_preset = PROVIDER_PRESETS.get(image_provider_type, {})
    if not image_use_global and image_preset.get("needs_key") and not image_api_key.strip():
        errors.append("Укажите API ключ для провайдера изображений или выберите «Использовать глобальные настройки».")
    if not image_use_global and image_provider_type == "custom" and not image_base_url.strip():
        errors.append("Для кастомного image-провайдера укажите Base URL.")
    if not image_preset.get("supports_images", True) and not image_use_global:
        errors.append(f"Провайдер «{image_preset.get('name')}» не поддерживает генерацию изображений. Выберите другой.")

    if errors:
        return render(
            request, "provider.html", status_code=400,
            user=user, provider_presets=PROVIDER_PRESETS, errors=errors,
            form={
                "text_provider_type": text_provider_type, "text_api_key": text_api_key,
                "text_base_url": text_base_url, "text_model": text_model,
                "image_provider_type": image_provider_type, "image_api_key": image_api_key,
                "image_base_url": image_base_url, "image_model": image_model,
            },
        )

    # Если "использовать глобальные" — сохраняем пустой ключ (будет использован глобальный)
    final_text_key = "" if text_use_global else text_api_key.strip()
    final_image_key = "" if image_use_global else image_api_key.strip()

    # Подставляем дефолтные значения из пресетов
    final_text_base = text_base_url.strip() or text_preset.get("base_url", "")
    final_text_model = text_model.strip() or text_preset.get("text_model", "")
    final_image_base = image_base_url.strip() or image_preset.get("base_url", "")
    final_image_model = image_model.strip() or image_preset.get("image_model", "")

    save_provider_settings(
        user["id"],
        text_provider_type, final_text_key, final_text_base, final_text_model,
        image_provider_type, final_image_key, final_image_base, final_image_model,
    )
    flash(request, "Провайдеры ИИ настроены: текст — «%s», изображения — «%s»." % (
        text_preset.get("name", text_provider_type),
        image_preset.get("name", image_provider_type),
    ))
    return RedirectResponse("/provider", status_code=303)


@app.get("/generate")
def generate_form(request: Request, user: dict = Depends(current_user),
                  plan_id: str = Query(""), platform: str = Query(""),
                  idea: str = Query("")):
    latest_survey = get_latest_brand_survey_for_user(user["id"])
    surveys = get_brand_surveys_for_user(user["id"])
    prefill = None
    if plan_id:
        prefill = {"plan_id": plan_id, "platform": platform, "idea": idea}
    return render(request, "generate.html", user=user, platforms=PLATFORMS,
                  latest_survey=latest_survey, surveys=surveys, prefill=prefill)


@app.post("/generate")
def generate_submit(
    request: Request,
    platform: str = Form(...),
    style: str = Form(...),
    wishes: str = Form(""),
    survey_id: str = Form(""),
    user: dict = Depends(current_user),
):
    # Выбираем опрос медиа-образа: явно указанный или последний
    brand_data = None
    if survey_id:
        try:
            sid = int(survey_id)
        except ValueError:
            sid = None
        if sid:
            survey = get_brand_survey_by_id(sid, user["id"])
            if survey:
                brand_data = survey["data"]
    if not brand_data:
        latest = get_latest_brand_survey_for_user(user["id"])
        if latest:
            brand_data = latest["data"]

    if not brand_data:
        # Нет медиа-образа — возвращаем форму с подсказкой
        surveys = get_brand_surveys_for_user(user["id"])
        return render(
            request, "generate.html", status_code=400,
            user=user, platforms=PLATFORMS, surveys=surveys,
            latest_survey=None,
            errors=["Сначала создайте медиа-образ — посты генерируются на его основе."],
        )

    user_gen = get_user_content_gen(user)
    try:
        content = user_gen.generate_post_from_brand_profile(brand_data, platform, style, wishes)
    except Exception as e:
        surveys = get_brand_surveys_for_user(user["id"])
        latest_survey = get_latest_brand_survey_for_user(user["id"])
        return render(
            request, "generate.html", status_code=502,
            user=user, platforms=PLATFORMS, surveys=surveys,
            latest_survey=latest_survey,
            errors=[f"Ошибка генерации: {e}. Проверьте настройки провайдера в разделе «ИИ»."],
            form={"platform": platform, "style": style, "wishes": wishes},
        )
    save_post_for_user(user["id"], platform, content, style, wishes)
    latest_survey = get_latest_brand_survey_for_user(user["id"])
    return render(
        request, "generate.html", user=user, platforms=PLATFORMS, result=content,
        latest_survey=latest_survey,
        form={"platform": platform, "style": style, "wishes": wishes},
    )

@app.get("/history")
def history(request: Request, user: dict = Depends(current_user)):
    posts = get_posts_for_user(user["id"], limit=50)
    # Добавляем изображения к каждому посту
    for p in posts:
        p["images"] = get_post_images(p["id"], user["id"])
    return render(request, "history.html", user=user, posts=posts, platform_label=platform_label)


@app.get("/posts/{post_id}")
def post_detail(request: Request, post_id: int, user: dict = Depends(current_user)):
    post = get_post_by_id_for_user(post_id, user["id"])
    if not post:
        raise HTTPException(status_code=404, detail="Пост не найден")
    images = get_post_images(post_id, user["id"])
    revisions = get_post_revisions(post_id, user["id"])
    return render(request, "post_detail.html", user=user, post=post,
                  images=images, revisions=revisions,
                  platform_label=platform_label, templates=IMAGE_TEMPLATES)


@app.get("/posts/{post_id}/edit")
def post_edit_form(request: Request, post_id: int, user: dict = Depends(current_user)):
    post = get_post_by_id_for_user(post_id, user["id"])
    if not post:
        raise HTTPException(status_code=404, detail="Пост не найден")
    return render(request, "post_edit.html", user=user, post=post, platform_label=platform_label)


@app.post("/posts/{post_id}/edit")
def post_edit_submit(
    request: Request,
    post_id: int,
    content: str = Form(...),
    user: dict = Depends(current_user),
):
    post = get_post_by_id_for_user(post_id, user["id"])
    if not post:
        raise HTTPException(status_code=404, detail="Пост не найден")
    content = content.strip()
    if not content:
        return render(
            request, "post_edit.html", status_code=400,
            user=user, post=post, platform_label=platform_label,
            errors=["Текст поста не может быть пустым."],
        )
    update_post_content(user["id"], post_id, content)
    add_post_revision(post_id, user["id"], content)
    flash(request, "Пост обновлён.")
    return RedirectResponse(f"/posts/{post_id}", status_code=303)


@app.post("/api/posts/{post_id}/edit-text")
def api_post_edit_text(
    post_id: int,
    content: str = Form(...),
    user: dict = Depends(current_user),
):
    """AJAX-сохранение текста поста (inline-редактор). Сохраняет ревизию."""
    post = get_post_by_id_for_user(post_id, user["id"])
    if not post:
        raise HTTPException(status_code=404, detail="Пост не найден")
    content = content.strip()
    if not content:
        return JSONResponse({"error": "Текст не может быть пустым"}, status_code=400)

    update_post_content(user["id"], post_id, content)
    add_post_revision(post_id, user["id"], content)

    return JSONResponse({"ok": True, "content": content})


@app.post("/posts/{post_id}/delete")
def post_delete(request: Request, post_id: int, user: dict = Depends(current_user)):
    post = get_post_by_id_for_user(post_id, user["id"])
    if not post:
        raise HTTPException(status_code=404, detail="Пост не найден")
    # Удаляем изображения поста
    images = get_post_images(post_id, user["id"])
    for img in images:
        fp = MEDIA_DIR / img["filename"]
        if fp.exists():
            fp.unlink()
    delete_post_by_id(user["id"], post_id)
    flash(request, "Пост удалён.")
    return RedirectResponse("/history", status_code=303)


# === ИЗОБРАЖЕНИЯ ПОСТА ===

@app.post("/posts/{post_id}/images/upload")
async def post_image_upload(request: Request, post_id: int, user: dict = Depends(current_user)):
    post = get_post_by_id_for_user(post_id, user["id"])
    if not post:
        raise HTTPException(status_code=404, detail="Пост не найден")

    form = await request.form()
    files = form.getlist("files")
    if not files:
        flash(request, "Выберите файлы для загрузки.")
        return RedirectResponse(f"/posts/{post_id}", status_code=303)

    uploaded = 0
    for upload in files:
        if not upload or not upload.filename:
            continue
        raw = await upload.read()
        if not raw:
            continue
        ext = detect_ext(raw)
        filename = f"{uuid.uuid4().hex}{ext}"
        (MEDIA_DIR / filename).write_bytes(raw)
        save_post_image(post_id, user["id"], filename, source="upload")
        uploaded += 1

    if uploaded:
        flash(request, f"Загружено изображений: {uploaded}")
    else:
        flash(request, "Не удалось загрузить изображения.")
    return RedirectResponse(f"/posts/{post_id}", status_code=303)


@app.post("/posts/{post_id}/images/generate")
def post_image_generate(request: Request, post_id: int, user: dict = Depends(current_user)):
    post = get_post_by_id_for_user(post_id, user["id"])
    if not post:
        raise HTTPException(status_code=404, detail="Пост не найден")

    # Генерируем промпт из текста поста
    user_img_gen = get_user_image_gen(user)
    # Берём первые 500 символов поста как основу для промпта
    post_text = (post["content"] or "")[:500]
    prompt = f"Create a professional social media image for this post. Post text: {post_text}. Style: modern, clean, eye-catching."

    try:
        image_bytes = user_img_gen.generate(prompt, "square")
    except Exception as e:
        flash(request, f"Ошибка генерации изображения: {e}")
        return RedirectResponse(f"/posts/{post_id}", status_code=303)

    ext = detect_ext(image_bytes)
    filename = f"{uuid.uuid4().hex}{ext}"
    (MEDIA_DIR / filename).write_bytes(image_bytes)
    save_post_image(post_id, user["id"], filename, source="generated")
    flash(request, "Изображение сгенерировано.")
    return RedirectResponse(f"/posts/{post_id}", status_code=303)


@app.post("/posts/{post_id}/images/{image_id}/delete")
def post_image_delete(request: Request, post_id: int, image_id: int, user: dict = Depends(current_user)):
    post = get_post_by_id_for_user(post_id, user["id"])
    if not post:
        raise HTTPException(status_code=404, detail="Пост не найден")
    filename = delete_post_image(image_id, user["id"])
    if filename:
        fp = MEDIA_DIR / filename
        if fp.exists():
            fp.unlink()
        flash(request, "Изображение удалено.")
    return RedirectResponse(f"/posts/{post_id}", status_code=303)


@app.get("/plan")
def plan_form(request: Request, user: dict = Depends(current_user)):
    items = get_plan_for_user(user["id"])
    status_counts = get_plan_stats_for_user(user["id"])
    return render(
        request, "plan.html",
        user=user, plan=items, status_counts=status_counts, platform_label=platform_label,
    )


@app.post("/plan")
def plan_submit(request: Request, days: int = Form(7), user: dict = Depends(current_user)):
    days = max(1, min(days, 14))
    user_gen = get_user_content_gen(user)
    latest = get_latest_brand_survey_for_user(user["id"])
    if latest:
        plan = user_gen.generate_content_plan_from_brand_profile(latest["data"], days)
    else:
        # Fallback — простой шаблонный план без медиа-образа
        plan = user_gen.generate_content_plan("бизнес", "не указано", days)
    replace_plan_for_user(user["id"], plan)
    flash(request, f"Контент-план на {days} дн. сохранён.")
    return RedirectResponse("/plan", status_code=303)


@app.post("/plan/{item_id}/status")
def plan_status_submit(
    request: Request,
    item_id: int,
    status: str = Form(...),
    user: dict = Depends(current_user),
):
    try:
        update_plan_status(user["id"], item_id, status)
        flash(request, "Статус пункта обновлён.")
    except ValueError:
        flash(request, "Недопустимый статус.")
    return RedirectResponse("/plan", status_code=303)


@app.post("/plan/{item_id}/delete")
def plan_item_delete(request: Request, item_id: int, user: dict = Depends(current_user)):
    delete_plan_item(user["id"], item_id)
    flash(request, "Пункт плана удалён.")
    return RedirectResponse("/plan", status_code=303)


@app.get("/stats")
def stats(request: Request, user: dict = Depends(current_user)):
    return render(request, "stats.html", user=user, stats=get_stats_for_user(user["id"]))


@app.get("/delete-account")
def delete_confirm_form(request: Request, user: dict = Depends(current_user)):
    return render(request, "delete_confirm.html", user=user)


@app.post("/delete-account")
def delete_account_submit(
    request: Request, confirm: str = Form(""), user: dict = Depends(current_user)
):
    if confirm != "yes":
        return RedirectResponse("/dashboard", status_code=303)
    delete_user_by_id(user["id"])
    request.session.clear()
    return RedirectResponse("/login", status_code=303)


# === ГЕНЕРАЦИЯ КАРТИНОК ===

@app.get("/images")
def images_gallery(request: Request, user: dict = Depends(current_user)):
    items = get_media_for_user(user["id"])
    return render(request, "images.html", user=user, images=items)


@app.get("/images/generate")
def image_generate_form(request: Request, user: dict = Depends(current_user)):
    # Размеры зависят от модели image-провайдера
    img_settings = get_image_provider_settings(user["id"])
    if img_settings and img_settings.get("image_model"):
        sizes = get_sizes_for_model(img_settings["image_model"])
    else:
        sizes = IMAGE_SIZES
    return render(
        request, "image_generate.html",
        user=user, sizes=sizes, styles=STYLE_PRESETS, templates=IMAGE_TEMPLATES,
    )


@app.post("/images/generate")
def image_generate_submit(
    request: Request,
    prompt: str = Form(...),
    size: str = Form("square"),
    style: str = Form("none"),
    user: dict = Depends(current_user),
):
    # Размеры зависят от модели
    img_settings = get_image_provider_settings(user["id"])
    if img_settings and img_settings.get("image_model"):
        sizes = get_sizes_for_model(img_settings["image_model"])
    else:
        sizes = IMAGE_SIZES

    prompt = prompt.strip()
    if not prompt:
        return render(
            request, "image_generate.html", status_code=400,
            user=user, sizes=sizes, styles=STYLE_PRESETS,
            form={"prompt": prompt, "size": size, "style": style},
            errors=["Опиши, что нужно нарисовать."],
        )

    user_img_gen = get_user_image_gen(user)
    full_prompt = user_img_gen.build_prompt(prompt, style)
    try:
        image_bytes = user_img_gen.generate(full_prompt, size)
    except Exception as e:
        return render(
            request, "image_generate.html", status_code=502,
            user=user, sizes=sizes, styles=STYLE_PRESETS,
            form={"prompt": prompt, "size": size, "style": style},
            errors=[f"Ошибка генерации картинки: {e}"],
        )

    filename = f"{uuid.uuid4().hex}{detect_ext(image_bytes)}"
    (MEDIA_DIR / filename).write_bytes(image_bytes)
    save_media(user["id"], prompt, style, filename)
    flash(request, "Картинка сгенерирована и сохранена в библиотеку.")
    return RedirectResponse("/images", status_code=303)


@app.post("/images/{media_id}/delete")
def image_delete(request: Request, media_id: int, user: dict = Depends(current_user)):
    item = get_media_by_id_for_user(media_id, user["id"])
    if not item:
        raise HTTPException(status_code=404, detail="Картинка не найдена")
    filename = delete_media_by_id(user["id"], media_id)
    if filename:
        file_path = MEDIA_DIR / filename
        if file_path.exists():
            file_path.unlink()
    flash(request, "Картинка удалена.")
    return RedirectResponse("/images", status_code=303)


# === ОПРОСЫ МЕДИА-ОБРАЗА ===

@app.get("/surveys")
def survey_list(request: Request, user: dict = Depends(current_user)):
    surveys = get_brand_surveys_for_user(user["id"])
    return render(request, "survey_list.html", user=user, surveys=surveys)


@app.get("/survey")
def survey_form(request: Request, user: dict = Depends(current_user)):
    return render(request, "survey.html", user=user)


@app.get("/survey/import")
def survey_import_form(request: Request, user: dict = Depends(current_user)):
    return render(request, "survey_import.html", user=user)


@app.post("/survey/import")
async def survey_import_submit(request: Request, user: dict = Depends(current_user)):
    import json as _json
    from docx import Document as _DocxDocument
    from io import BytesIO as _BytesIO

    form = await request.form()
    upload = form.get("file")
    if not upload or not upload.filename:
        return render(
            request, "survey_import.html", status_code=400,
            user=user, errors=["Выберите файл для загрузки."],
        )

    filename = upload.filename.lower()
    raw = await upload.read()

    # --- JSON: прямой импорт ---
    if filename.endswith(".json"):
        try:
            data = _json.loads(raw.decode("utf-8"))
        except Exception as e:
            return render(
                request, "survey_import.html", status_code=400,
                user=user, errors=[f"Ошибка чтения JSON: {e}"],
            )
        # Поддерживаем два формата: {"data": {...}, "title": "..."} или просто {...}
        if isinstance(data, dict) and "data" in data and isinstance(data["data"], dict):
            title = (data.get("title") or "Импортированный образ").strip()[:200]
            survey_data = data["data"]
        elif isinstance(data, dict):
            title = (data.get("client_name") or data.get("title") or "Импортированный образ").strip()[:200]
            survey_data = data
        else:
            return render(
                request, "survey_import.html", status_code=400,
                user=user, errors=["JSON должен быть объектом с полями опроса."],
            )
        # Гарантируем структуры
        survey_data.setdefault("channels", [])
        survey_data.setdefault("funnel", [])
        survey_id = save_brand_survey(user["id"], title, survey_data)
        flash(request, f"Медиа-образ «{title}» импортирован из JSON.")
        return RedirectResponse(f"/surveys/{survey_id}", status_code=303)

    # --- DOCX: разбор через PiAPI ---
    if filename.endswith(".docx"):
        try:
            doc = _DocxDocument(_BytesIO(raw))
        except Exception as e:
            return render(
                request, "survey_import.html", status_code=400,
                user=user, errors=[f"Ошибка чтения .docx: {e}"],
            )

        # Извлекаем весь текст
        lines = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                lines.append(text)
        # Текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                lines.append(" | ".join(cells))
        doc_text = "\n".join(lines)

        if not doc_text.strip():
            return render(
                request, "survey_import.html", status_code=400,
                user=user, errors=["Документ пустой — нет текста для разбора."],
            )

        # Разбор через PiAPI
        survey_data = content_gen.parse_brand_doc_to_survey(doc_text)
        if not survey_data:
            return render(
                request, "survey_import.html", status_code=502,
                user=user,
                errors=["Не удалось разобрать документ. Проверьте, что PiAPI ключ настроен, и попробуйте снова."],
            )

        title = (survey_data.get("client_name") or "Импортированный образ").strip()[:200]
        survey_id = save_brand_survey(user["id"], title, survey_data)
        flash(request, f"Медиа-образ «{title}» импортирован из .docx.")
        return RedirectResponse(f"/surveys/{survey_id}", status_code=303)

    return render(
        request, "survey_import.html", status_code=400,
        user=user, errors=["Поддерживаются только файлы .docx и .json"],
    )


@app.post("/survey/submit")
async def survey_submit(request: Request, user: dict = Depends(current_user)):
    body = await request.json()
    title = (body.get("title") or "Без названия").strip()[:200]
    data = body.get("data") or {}
    survey_id = save_brand_survey(user["id"], title, data)
    return {"id": survey_id, "ok": True}


@app.get("/surveys/{survey_id}")
def survey_view(request: Request, survey_id: int, user: dict = Depends(current_user)):
    survey = get_brand_survey_by_id(survey_id, user["id"])
    if not survey:
        raise HTTPException(status_code=404, detail="Опрос не найден")
    return render(request, "survey_view.html", user=user, survey=survey)


@app.get("/surveys/{survey_id}/export")
def survey_export(request: Request, survey_id: int, user: dict = Depends(current_user)):
    survey = get_brand_survey_by_id(survey_id, user["id"])
    if not survey:
        raise HTTPException(status_code=404, detail="Опрос не найден")
    docx_bytes = generate_brand_docx(survey["data"])
    # ASCII-имя для Content-Disposition (latin-1), кириллица не допускается в заголовках
    safe_title = "".join(c if c.isascii() and c.isalnum() else "_" for c in survey["title"])[:80]
    if not safe_title or safe_title == "_" * len(safe_title):
        safe_title = "brand"
    filename = f"Strategiya_brenda_{safe_title}.docx"
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.post("/surveys/{survey_id}/delete")
def survey_delete(request: Request, survey_id: int, user: dict = Depends(current_user)):
    survey = get_brand_survey_by_id(survey_id, user["id"])
    if not survey:
        raise HTTPException(status_code=404, detail="Опрос не найден")
    delete_brand_survey(user["id"], survey_id)
    flash(request, "Опрос удалён.")
    return RedirectResponse("/surveys", status_code=303)


# === КАНАЛЫ ДЛЯ ПУБЛИКАЦИИ ===

@app.get("/channels")
def channels_list(request: Request, user: dict = Depends(current_user)):
    channels = get_channels_for_user(user["id"])
    return render(request, "channels.html", user=user, channels=channels)


@app.get("/channels/add")
def channel_add_form(request: Request, user: dict = Depends(current_user)):
    return render(request, "channel_add.html", user=user)


@app.post("/channels/add")
def channel_add_submit(
    request: Request,
    chat_input: str = Form(...),
    user: dict = Depends(current_user),
):
    chat_input = chat_input.strip()
    if not chat_input:
        return render(
            request, "channel_add.html", status_code=400,
            user=user, errors=["Введите @username или ID канала/чата."],
            form={"chat_input": chat_input},
        )

    # 1. Разрешаем чат через Telegram API
    info = resolve_chat(chat_input)
    if "error" in info:
        return render(
            request, "channel_add.html", status_code=400,
            user=user, errors=[f"Не удалось найти чат: {info['error']}"],
            form={"chat_input": chat_input},
        )

    chat_id = info["chat_id"]
    title = info["title"]
    username = info["username"]
    chat_type = info["type"]

    # 2. Проверяем, что бот — админ
    is_admin, admin_error = check_bot_is_admin(chat_id)
    if not is_admin:
        return render(
            request, "channel_add.html", status_code=400,
            user=user, errors=[admin_error or "Бот не админ в этом чате."],
            form={"chat_input": chat_input},
        )

    # 3. Проверяем, не добавлен ли уже
    if is_channel_already_added(user["id"], chat_id):
        return render(
            request, "channel_add.html", status_code=400,
            user=user, errors=[f"Канал «{title}» уже добавлен."],
            form={"chat_input": chat_input},
        )

    # 4. Сохраняем
    save_channel(user["id"], chat_id, title, username, chat_type)
    flash(request, f"Канал «{title}» добавлен.")
    return RedirectResponse("/channels", status_code=303)


@app.post("/channels/{channel_id}/delete")
def channel_delete(request: Request, channel_id: int, user: dict = Depends(current_user)):
    channel = get_channel_by_id(channel_id, user["id"])
    if not channel:
        raise HTTPException(status_code=404, detail="Канал не найден")
    delete_channel(user["id"], channel_id)
    flash(request, f"Канал «{channel['title']}» удалён.")
    return RedirectResponse("/channels", status_code=303)


# === ПУБЛИКАЦИЯ ПОСТА ===

@app.get("/posts/{post_id}/publish")
def publish_form(request: Request, post_id: int, user: dict = Depends(current_user)):
    post = get_post_by_id_for_user(post_id, user["id"])
    if not post:
        raise HTTPException(status_code=404, detail="Пост не найден")
    channels = get_channels_for_user(user["id"])
    return render(request, "publish.html", user=user, post=post, channels=channels,
                  platform_label=platform_label)


@app.post("/posts/{post_id}/publish")
def publish_submit(
    request: Request,
    post_id: int,
    channel_id: str = Form(...),
    user: dict = Depends(current_user),
):
    post = get_post_by_id_for_user(post_id, user["id"])
    if not post:
        raise HTTPException(status_code=404, detail="Пост не найден")

    try:
        cid = int(channel_id)
    except ValueError:
        cid = None
    if not cid:
        channels = get_channels_for_user(user["id"])
        return render(
            request, "publish.html", status_code=400,
            user=user, post=post, channels=channels, platform_label=platform_label,
            errors=["Выберите канал для публикации."],
        )

    channel = get_channel_by_id(cid, user["id"])
    if not channel:
        raise HTTPException(status_code=404, detail="Канал не найден")

    # Берём первую картинку поста (если есть) для публикации в Telegram
    image_bytes = None
    images = get_post_images(post_id, user["id"])
    if images:
        fp = MEDIA_DIR / images[0]["filename"]
        if fp.exists():
            image_bytes = fp.read_bytes()

    ok, message = publish_post(channel["chat_id"], post["content"], image_bytes=image_bytes)
    if ok:
        mark_post_published(post_id, user["id"], channel["title"])
        flash(request, f"✅ Пост опубликован в «{channel['title']}»")
    else:
        flash(request, f"❌ Ошибка публикации: {message}")

    return RedirectResponse(f"/posts/{post_id}", status_code=303)


# === ФОНОВАЯ ГЕНЕРАЦИЯ (AJAX) ===

def _run_post_generation(task_id, user_id, platform, style, wishes, survey_id, plan_id=None):
    """Фоновая генерация поста."""
    try:
        update_task_status(task_id, 'running')
        brand_data = None
        if survey_id:
            survey = get_brand_survey_by_id(survey_id, user_id)
            if survey:
                brand_data = survey["data"]
        if not brand_data:
            latest = get_latest_brand_survey_for_user(user_id)
            if latest:
                brand_data = latest["data"]
        if not brand_data:
            update_task_status(task_id, 'error', error="Нет медиа-образа. Создайте медиа-образ и повторите.")
            return
        settings = get_text_provider_settings(user_id)
        gen = ContentGenerator(provider_settings=settings) if settings and settings.get("api_key") else content_gen
        content = gen.generate_post_from_brand_profile(brand_data, platform, style, wishes)
        post_id = save_post_for_user(user_id, platform, content, style, wishes)
        # Если пост создан из плана — ставим статус «опубликовано»
        if plan_id:
            try:
                update_plan_status(user_id, int(plan_id), "published")
            except (ValueError, Exception) as e:
                print(f"⚠️ Не удалось обновить статус плана {plan_id}: {e}")
        update_task_status(task_id, 'done', result_json=json.dumps({"post_id": post_id, "content": content}))
    except Exception as e:
        update_task_status(task_id, 'error', error=str(e))


def _run_plan_generation(task_id, user_id, days):
    """Фоновая генерация контент-плана."""
    try:
        update_task_status(task_id, 'running')
        settings = get_text_provider_settings(user_id)
        gen = ContentGenerator(provider_settings=settings) if settings and settings.get("api_key") else content_gen
        latest = get_latest_brand_survey_for_user(user_id)
        if latest:
            plan = gen.generate_content_plan_from_brand_profile(latest["data"], days)
        else:
            plan = gen.generate_content_plan("бизнес", "не указано", days)
        replace_plan_for_user(user_id, plan)
        update_task_status(task_id, 'done', result_json=json.dumps({"days": days, "items": len(plan)}))
    except Exception as e:
        update_task_status(task_id, 'error', error=str(e))


def _run_modify_post(task_id, user_id, post_id, instruction):
    """Фоновая модификация поста по инструкции пользователя."""
    try:
        update_task_status(task_id, 'running')

        post = get_post_by_id_for_user(post_id, user_id)
        if not post:
            update_task_status(task_id, 'error', error="Пост не найден")
            return

        if not instruction.strip():
            update_task_status(task_id, 'error', error="Инструкция не может быть пустой")
            return

        # Получаем AI-провайдер
        settings = get_text_provider_settings(user_id)
        gen = ContentGenerator(provider_settings=settings) if settings and settings.get("api_key") else content_gen

        new_content = gen.modify_post(post["content"], instruction.strip())

        # Сохраняем новую версию
        update_post_content(user_id, post_id, new_content)
        add_post_revision(post_id, user_id, new_content)

        update_task_status(task_id, 'done', result_json=json.dumps({
            "content": new_content,
            "instruction": instruction.strip(),
        }, ensure_ascii=False))

    except Exception as e:
        update_task_status(task_id, 'error', error=str(e))


def _run_social_analysis(task_id, user_id, url, manual_posts):
    """Фоновый анализ профиля соцсети."""
    try:
        update_task_status(task_id, 'running')

        platform_info = detect_platform(url)
        posts = []

        # Если пользователь вставил тексты постов вручную
        if manual_posts and manual_posts.strip():
            blocks = re.split(r'\n{2,}', manual_posts.strip())
            for block in blocks:
                block = block.strip()
                if len(block) > 20:
                    posts.append({"text": block, "date": ""})

        # Если URL — Telegram, загружаем посты автоматически
        if not posts and platform_info["platform"] == "telegram" and platform_info["fetchable"]:
            posts = fetch_telegram_posts(platform_info["channel"])

        if not posts:
            update_task_status(task_id, 'error',
                error="Не удалось загрузить посты. Для VK и Tenchat вставьте тексты постов вручную (по 2+ постов через пустую строку).")
            return

        # AI-анализ
        settings = get_text_provider_settings(user_id)
        result = analyze_posts_with_ai(posts, provider_settings=settings)

        if "error" in result:
            update_task_status(task_id, 'error', error=result["error"])
            return

        result["platform_info"] = platform_info
        result["posts_count"] = len(posts)
        update_task_status(task_id, 'done', result_json=json.dumps(result, ensure_ascii=False))

    except Exception as e:
        update_task_status(task_id, 'error', error=str(e))


def _run_image_generation(task_id, user_id, prompt, size, style, reference_image_bytes=None):
    """Фоновая генерация картинки."""
    try:
        update_task_status(task_id, 'running')
        settings = get_image_provider_settings(user_id)
        gen = ImageGenerator(provider_settings=settings) if settings and settings.get("api_key") else image_gen
        full_prompt = gen.build_prompt(prompt, style)
        if reference_image_bytes:
            # Для image-to-image нужен chat/completions — используем text-провайдер
            text_settings = get_text_provider_settings(user_id)
            image_bytes = gen.generate_with_reference(full_prompt, reference_image_bytes, size, chat_settings=text_settings)
        else:
            image_bytes = gen.generate(full_prompt, size)
        filename = f"{uuid.uuid4().hex}{detect_ext(image_bytes)}"
        (MEDIA_DIR / filename).write_bytes(image_bytes)
        save_media(user_id, prompt, style, filename)
        update_task_status(task_id, 'done', result_json=json.dumps({"filename": filename}))
    except Exception as e:
        update_task_status(task_id, 'error', error=str(e))


def _run_analyze_post_image(task_id, user_id, post_id):
    """Фоновый анализ поста — предлагает концепции изображений."""
    try:
        update_task_status(task_id, 'running')
        post = get_post_by_id_for_user(post_id, user_id)
        if not post:
            update_task_status(task_id, 'error', error="Пост не найден")
            return
        post_text = (post["content"] or "")[:2000]

        # Получаем медиа-образ для контекста бренда
        brand_data = None
        latest = get_latest_brand_survey_for_user(user_id)
        if latest:
            brand_data = latest["data"]

        settings = get_text_provider_settings(user_id)
        gen = ContentGenerator(provider_settings=settings) if settings and settings.get("api_key") else content_gen
        concepts = gen.suggest_image_concepts(post_text, brand_data)
        update_task_status(task_id, 'done', result_json=json.dumps({"concepts": concepts}))
    except Exception as e:
        update_task_status(task_id, 'error', error=str(e))


def _run_post_image_generation(task_id, user_id, post_id, template_key=None, custom_prompt=None, custom_style=None, custom_size=None):
    """Фоновая генерация картинки к посту."""
    try:
        update_task_status(task_id, 'running')
        post = get_post_by_id_for_user(post_id, user_id)
        if not post:
            update_task_status(task_id, 'error', error="Пост не найден")
            return
        settings = get_image_provider_settings(user_id)
        gen = ImageGenerator(provider_settings=settings) if settings and settings.get("api_key") else image_gen

        post_text = (post["content"] or "")[:500]

        if custom_prompt:
            # Концепт от ИИ — используем его промпт + текст поста
            prompt = f"{custom_prompt}. Post text for context: {post_text}"
            size = custom_size or "wide"
            style = custom_style or "none"
        elif template_key:
            # Ищем шаблон по ключу
            tpl = None
            for t in IMAGE_TEMPLATES:
                if t["key"] == template_key:
                    tpl = t
                    break
            if tpl:
                prompt = tpl["prompt_template"].replace("{post_text}", post_text)
                size = tpl["size"]
                style = tpl["style"]
            else:
                prompt = f"Create a professional social media image for this post. Post text: {post_text}. Style: modern, clean, eye-catching."
                size = "square"
                style = None
        else:
            # Старое поведение — авто-промпт
            prompt = f"Create a professional social media image for this post. Post text: {post_text}. Style: modern, clean, eye-catching."
            size = "square"
            style = None

        if style:
            prompt = gen.build_prompt(prompt, style)

        image_bytes = gen.generate(prompt, size)
        ext = detect_ext(image_bytes)
        filename = f"{uuid.uuid4().hex}{ext}"
        (MEDIA_DIR / filename).write_bytes(image_bytes)
        save_post_image(post_id, user_id, filename, source="generated")
        update_task_status(task_id, 'done', result_json=json.dumps({"filename": filename}))
    except Exception as e:
        update_task_status(task_id, 'error', error=str(e))


@app.post("/api/tasks")
async def api_start_task(request: Request, user: dict = Depends(current_user)):
    """Запускает фоновую задачу. Возвращает task_id."""
    uid = user["id"]
    form = await request.form()
    task_type_val = form.get("task_type", "")
    task_id = create_task(uid, task_type_val)

    if task_type_val == "generate_post":
        platform = form.get("platform", "telegram")
        style = form.get("style", "")
        wishes = form.get("wishes", "")
        survey_id = form.get("survey_id", "")
        plan_id = form.get("plan_id", "") or None
        t = threading.Thread(target=_run_post_generation, args=(task_id, uid, platform, style, wishes, survey_id, plan_id), daemon=True)
        t.start()

    elif task_type_val == "generate_plan":
        days = int(form.get("days", 7))
        days = max(1, min(days, 14))
        t = threading.Thread(target=_run_plan_generation, args=(task_id, uid, days), daemon=True)
        t.start()

    elif task_type_val == "generate_image":
        prompt = form.get("prompt", "")
        size = form.get("size", "square")
        style = form.get("style", "none")
        ref_file = form.get("reference_image")
        ref_bytes = None
        if ref_file and hasattr(ref_file, "read"):
            ref_bytes = await ref_file.read()
            if len(ref_bytes) < 100:
                ref_bytes = None  # слишком маленький, пропускаем
        t = threading.Thread(target=_run_image_generation, args=(task_id, uid, prompt, size, style, ref_bytes), daemon=True)
        t.start()

    elif task_type_val == "generate_post_image":
        post_id = int(form.get("post_id", 0))
        template_key = form.get("template_key", "") or None
        t = threading.Thread(target=_run_post_image_generation, args=(task_id, uid, post_id, template_key), daemon=True)
        t.start()

    elif task_type_val == "analyze_post_image":
        post_id = int(form.get("post_id", 0))
        t = threading.Thread(target=_run_analyze_post_image, args=(task_id, uid, post_id), daemon=True)
        t.start()

    elif task_type_val == "generate_post_image_with_concept":
        post_id = int(form.get("post_id", 0))
        concept_prompt = form.get("concept_prompt", "")
        concept_style = form.get("concept_style", "none")
        concept_size = form.get("concept_size", "square")
        t = threading.Thread(
            target=_run_post_image_generation,
            args=(task_id, uid, post_id),
            kwargs={"custom_prompt": concept_prompt, "custom_style": concept_style, "custom_size": concept_size},
            daemon=True,
        )
        t.start()

    elif task_type_val == "analyze_social_profile":
        url = form.get("url", "")
        manual_posts = form.get("manual_posts", "")
        t = threading.Thread(
            target=_run_social_analysis,
            args=(task_id, uid, url, manual_posts),
            daemon=True,
        )
        t.start()

    elif task_type_val == "modify_post":
        post_id = int(form.get("post_id", 0))
        instruction = form.get("instruction", "")
        t = threading.Thread(
            target=_run_modify_post,
            args=(task_id, uid, post_id, instruction),
            daemon=True,
        )
        t.start()

    return {"task_id": task_id, "status": "pending"}


@app.get("/api/tasks/{task_id}")
def api_task_status(task_id: int, user: dict = Depends(current_user)):
    """Возвращает статус задачи."""
    task = get_task(task_id)
    if not task or task["user_id"] != user["id"]:
        raise HTTPException(status_code=404, detail="Задача не найдена")
    result = {
        "task_id": task["id"],
        "status": task["status"],
        "task_type": task["task_type"],
    }
    if task["status"] == "done" and task["result_json"]:
        result["result"] = json.loads(task["result_json"])
    if task["error"]:
        result["error"] = task["error"]
    return result
