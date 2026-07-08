from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

def create_report():
    # Создаём документ
    doc = Document()
    
    # Заголовок
    title = doc.add_heading('Отчет о разработке Telegram-бота для предпринимателей', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Дата
    doc.add_paragraph(f'Дата: {datetime.now().strftime("%d.%m.%Y")}')
    doc.add_paragraph()
    
    # 1. Введение
    doc.add_heading('1. Введение', level=1)
    doc.add_heading('1.1. Назначение приложения', level=2)
    doc.add_paragraph(
        'Данный Telegram-бот предназначен для помощи предпринимателям в ведении социальных сетей. '
        'Бот собирает информацию о пользователе, его бизнесе и на основе этих данных генерирует уникальный '
        'контент для различных социальных платформ (Instagram, LinkedIn, Telegram, Twitter/X).'
    )
    
    doc.add_heading('1.2. Основные функции', level=2)
    doc.add_paragraph('• Регистрация пользователей с сохранением данных в SQLite', style='List Bullet')
    doc.add_paragraph('• Выбор из 9 видов предпринимательской деятельности', style='List Bullet')
    doc.add_paragraph('• Генерация контента для соцсетей с использованием нейросетей', style='List Bullet')
    doc.add_paragraph('• Управление ботом через команды в Telegram', style='List Bullet')
    doc.add_paragraph('• Интуитивно понятный интерфейс с кнопками', style='List Bullet')
    
    # 2. Архитектура
    doc.add_heading('2. Архитектура приложения', level=1)
    doc.add_heading('2.1. Структура проекта', level=2)
    
    doc.add_paragraph('''
    telegram_bot/
    ├── .env                     # Переменные окружения
    ├── config.py                # Загрузка конфигурации
    ├── database.py              # Работа с SQLite
    ├── main.py                  # Точка входа
    ├── requirements.txt         # Список библиотек
    ├── ai_core/
    │   ├── __init__.py
    │   └── neural_network.py    # Нейросеть
    ├── handlers/
    │   ├── __init__.py
    │   ├── registration.py
    │   ├── menu.py
    │   └── content.py
    └── keyboards/
        ├── __init__.py
        └── reply_keyboards.py
    ''', style='Normal')
    
    # 3. Библиотеки
    doc.add_heading('3. Используемые библиотеки', level=1)
    doc.add_heading('3.1. Основные зависимости', level=2)
    
    libs = [
        ('pytelegrambotapi', 'Работа с Telegram Bot API'),
        ('sqlite3', 'Легковесная база данных (встроенная)'),
        ('python-dotenv', 'Загрузка переменных из .env'),
        ('requests', 'HTTP-запросы к API'),
        ('groq', 'API для работы с нейросетями Groq'),
        ('python-docx', 'Создание Word-документов'),
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Библиотека'
    hdr_cells[1].text = 'Назначение'
    
    for lib, desc in libs:
        row_cells = table.add_row().cells
        row_cells[0].text = lib
        row_cells[1].text = desc
    
    doc.add_paragraph()
    
    # 4. База данных
    doc.add_heading('4. База данных', level=1)
    doc.add_heading('4.1. Схема таблиц', level=2)
    
    doc.add_paragraph('Таблица users:')
    table2 = doc.add_table(rows=1, cols=3)
    table2.style = 'Table Grid'
    hdr = table2.rows[0].cells
    hdr[0].text = 'Поле'
    hdr[1].text = 'Тип'
    hdr[2].text = 'Описание'
    
    fields = [
        ('id', 'INTEGER PRIMARY KEY', 'Уникальный ID записи'),
        ('telegram_id', 'INTEGER UNIQUE', 'ID пользователя в Telegram'),
        ('name', 'TEXT', 'Имя пользователя'),
        ('phone', 'TEXT', 'Номер телефона'),
        ('business_type', 'INTEGER', 'ID вида бизнеса (1-9)'),
        ('description', 'TEXT', 'Описание бизнеса'),
        ('registered', 'BOOLEAN', 'Флаг завершения регистрации'),
    ]
    
    for field, typ, desc in fields:
        row = table2.add_row().cells
        row[0].text = field
        row[1].text = typ
        row[2].text = desc
    
    # 5. Виды бизнеса
    doc.add_heading('4.2. 9 видов предпринимательства', level=2)
    
    business_types = [
        ('1', '🛒 Розничная торговля'),
        ('2', '🏭 Производство'),
        ('3', '💻 IT-услуги'),
        ('4', '📚 Образование'),
        ('5', '🏥 Медицина'),
        ('6', '🚚 Логистика'),
        ('7', '🏨 Гостиничный бизнес'),
        ('8', '🎨 Маркетинг и реклама'),
        ('9', '🔧 Услуги (ремонт, клининг)'),
    ]
    
    table3 = doc.add_table(rows=1, cols=2)
    table3.style = 'Table Grid'
    hdr = table3.rows[0].cells
    hdr[0].text = 'ID'
    hdr[1].text = 'Название'
    
    for bid, bname in business_types:
        row = table3.add_row().cells
        row[0].text = bid
        row[1].text = bname
    
    # 6. Управление
    doc.add_heading('5. Управление ботом', level=1)
    doc.add_heading('5.1. Команды Telegram', level=2)
    
    table4 = doc.add_table(rows=1, cols=2)
    table4.style = 'Table Grid'
    hdr = table4.rows[0].cells
    hdr[0].text = 'Команда'
    hdr[1].text = 'Описание'
    
    commands = [
        ('/start', 'Начать работу / открыть меню'),
        ('/help', 'Показать справку'),
        ('/status', 'Показать статус бота'),
        ('/stop', 'Остановить бота'),
        ('/restart', 'Перезапустить бота'),
    ]
    
    for cmd, desc in commands:
        row = table4.add_row().cells
        row[0].text = cmd
        row[1].text = desc
    
    # 7. Нейросеть
    doc.add_heading('6. Нейросетевая составляющая', level=1)
    doc.add_heading('6.1. Используемые модели', level=2)
    
    table5 = doc.add_table(rows=1, cols=3)
    table5.style = 'Table Grid'
    hdr = table5.rows[0].cells
    hdr[0].text = 'Модель'
    hdr[1].text = 'Провайдер'
    hdr[2].text = 'Назначение'
    
    models = [
        ('llama-3.3-70b-versatile', 'Groq', 'Генерация постов'),
        ('mixtral-8x7b-32768', 'Groq', 'Альтернативная модель'),
    ]
    
    for model, provider, purpose in models:
        row = table5.add_row().cells
        row[0].text = model
        row[1].text = provider
        row[2].text = purpose
    
    # 8. Установка
    doc.add_heading('7. Установка и запуск', level=1)
    
    doc.add_heading('7.1. Установка зависимостей', level=2)
    doc.add_paragraph('pip install -r requirements.txt', style='Normal')
    
    doc.add_heading('7.2. Настройка переменных окружения', level=2)
    doc.add_paragraph('.env:')
    doc.add_paragraph('TELEGRAM_TOKEN=ваш_токен_от_BotFather', style='Normal')
    doc.add_paragraph('GROQ_API_KEY=gsk_ваш_ключ_из_console.groq.com', style='Normal')
    
    doc.add_heading('7.3. Запуск', level=2)
    doc.add_paragraph('python main.py', style='Normal')
    
    # 9. Примеры
    doc.add_heading('8. Примеры работы', level=1)
    doc.add_paragraph('Сценарий регистрации:')
    doc.add_paragraph('1. Пользователь вводит /start', style='List Number')
    doc.add_paragraph('2. Бот запрашивает имя', style='List Number')
    doc.add_paragraph('3. Бот запрашивает телефон (кнопка)', style='List Number')
    doc.add_paragraph('4. Бот показывает 9 видов бизнеса (кнопки)', style='List Number')
    doc.add_paragraph('5. Бот запрашивает описание бизнеса', style='List Number')
    doc.add_paragraph('6. Регистрация завершена → главное меню', style='List Number')
    
    # 10. Заключение
    doc.add_heading('9. Заключение', level=1)
    doc.add_paragraph(
        'Разработанный Telegram-бот представляет собой полностью функциональное решение '
        'для предпринимателей, желающих автоматизировать создание контента для социальных сетей.'
    )
    
    doc.add_heading('Ключевые достижения:', level=2)
    doc.add_paragraph('• Модульная архитектура', style='List Bullet')
    doc.add_paragraph('• Интеграция с нейросетями (Groq)', style='List Bullet')
    doc.add_paragraph('• Удобный интерфейс с кнопками', style='List Bullet')
    doc.add_paragraph('• Локальное хранение данных (SQLite)', style='List Bullet')
    doc.add_paragraph('• Гибкая система шаблонов', style='List Bullet')
    doc.add_paragraph('• Управление через Telegram-команды', style='List Bullet')
    
    # Сохраняем
    filename = f'Отчет_о_разработке_бота_{datetime.now().strftime("%Y%m%d")}.docx'
    doc.save(filename)
    print(f'✅ Отчет создан: {filename}')
    print(f'📁 Файл сохранен в: {os.path.abspath(filename)}')

if __name__ == '__main__':
    create_report()